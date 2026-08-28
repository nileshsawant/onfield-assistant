"""Ingest a user's own data into a per-user private RAG store.

Backs ``ofa --add-private`` / ``--list-private`` / ``--forget-private``.

The private store is a ChromaDB instance under ``$OFA_SCRATCH`` — separate
from the shared corpora in ``$OFA_VECTORDB`` — so a user can index their own
data and have it retrieved alongside the built-in collections without any
write access to the shared install. See ``retrieve_private_context`` in
``ofa_main.py`` for the read side.

Chunking / embedding reuse ``rebuild_indices.py`` so the private store is
byte-for-byte compatible with what the shared pipeline produces.
"""
from __future__ import annotations

import json
import os
import sys
from pathlib import Path

# Reuse the shared pipeline's helpers rather than re-implementing chunking.
_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))

import rebuild_indices as _ri  # noqa: E402

# Resolved lazily from ofa_main so the two modules cannot disagree on where
# the private store lives.
try:
    from ofa_main import (
        PRIVATE_VECTORDB_PATH,
        PRIVATE_SOURCES_PATH,
        OFA_SCRATCH,
    )
except Exception:  # pragma: no cover - defensive, ofa_main should import fine
    OFA_SCRATCH = os.environ.get("OFA_SCRATCH", os.path.expanduser("~/.local/state/ofa"))
    PRIVATE_VECTORDB_PATH = os.environ.get(
        "OFA_PRIVATE_VECTORDB", os.path.join(OFA_SCRATCH, "vectordb-private")
    )
    PRIVATE_SOURCES_PATH = os.path.join(OFA_SCRATCH, ".ofa_private_sources.json")

# Formats handled without any extra dependencies. .docx/.xlsx are a
# deliberate later addition (they need python-docx / openpyxl).
CODE_EXTENSIONS = [
    ".md", ".rst", ".txt", ".py", ".c", ".h", ".cpp", ".hpp", ".cc", ".cxx",
    ".f", ".f90", ".F90", ".jl", ".m", ".sh", ".bash", ".toml", ".yaml",
    ".yml", ".json", ".cfg", ".ini", ".tex", ".ipynb", ".inp", ".i",
]

# Office formats extracted to plain text (lossy — see extractors below).
# Handled separately from CODE_EXTENSIONS because they are binary and need
# python-docx / openpyxl rather than a plain read_text().
OFFICE_EXTENSIONS = [".docx", ".xlsx"]


def _err(msg: str) -> None:
    print(msg, file=sys.stderr)


def _extract_docx(path: Path) -> str:
    """Flatten a .docx to text: paragraphs, then each table row as
    tab-joined cells. Loses styling, comments, tracked changes, and
    embedded objects — enough for retrieval, not a faithful render."""
    from docx import Document
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _extract_xlsx(path: Path) -> str:
    """Flatten a .xlsx to text, one line per non-empty row, prefixed by
    sheet name. Reads computed values, not formulas (data_only=True), so a
    workbook never opened in Excel may show blanks for formula cells."""
    from openpyxl import load_workbook
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                lines.append(f"[{ws.title}] " + "\t".join(cells))
    wb.close()
    return "\n".join(lines)


def _process_office_file(path: Path, root: Path, collection: str) -> list:
    """Return (chunk_id, doc, metadata) triples for a .docx/.xlsx file,
    matching rebuild_indices.process_code_file's shape."""
    ext = path.suffix.lower()
    try:
        text = _extract_docx(path) if ext == ".docx" else _extract_xlsx(path)
    except Exception as e:
        _err(f"  [!] {path}: {ext} extraction failed ({e}); skipped")
        return []
    if not text.strip():
        return []
    relpath = path.relative_to(root).as_posix()
    prefix = f"[{root.name} {ext.lstrip('.')} - {relpath}]"
    out = []
    for i, chunk in enumerate(_ri.chunk_text(text, _ri.CODE_CHUNK, _ri.CODE_OVERLAP)):
        out.append((
            _ri.stable_id(collection, str(path), i),
            f"{prefix}\n{chunk}",
            {
                "source_type": "office",
                "source_root": root.name,
                "filepath": relpath,
                "chunk_index": i,
            },
        ))
    return out


def _walk_office(root: Path):
    """Yield .docx/.xlsx files under root, skipping the same noise dirs as
    the shared walkers and Office lock files (~$foo.docx)."""
    ext_set = {e.lower() for e in OFFICE_EXTENSIONS}
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in _ri._SKIP_DIRS and not d.startswith(".")]
        for fn in filenames:
            if fn.startswith("~$"):
                continue
            p = Path(dirpath) / fn
            if p.suffix.lower() in ext_set:
                yield p


def _secure_mkdir(path: str, mode: int = 0o700) -> None:
    """Create *path* (and parents) and force *mode* on the leaf.

    We chmod explicitly rather than trusting the umask or the enclosing
    scratch directory's mode, which varies by site — private data must not
    land in a world-readable directory.
    """
    os.makedirs(path, exist_ok=True)
    try:
        os.chmod(path, mode)
    except OSError as e:
        _err(f"Warning: could not chmod {path} to {oct(mode)}: {e}")


def _load_sources() -> dict:
    try:
        with open(PRIVATE_SOURCES_PATH) as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _save_sources(sources: dict) -> None:
    _secure_mkdir(os.path.dirname(PRIVATE_SOURCES_PATH) or ".", 0o700)
    fd = os.open(PRIVATE_SOURCES_PATH, os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
    try:
        os.write(fd, json.dumps(sources, indent=2).encode("utf-8") + b"\n")
    finally:
        os.close(fd)


def _client():
    import chromadb
    _secure_mkdir(PRIVATE_VECTORDB_PATH, 0o700)
    return chromadb.PersistentClient(path=PRIVATE_VECTORDB_PATH)


def _embedder():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer(str(_ri.EMBEDDING_MODEL_PATH), device="cpu")


def _sanitize_collection_name(raw: str) -> str:
    """ChromaDB names must be 3-63 chars, [a-zA-Z0-9._-], start/end
    alphanumeric. Map an arbitrary label onto that, prefixed so private
    collections are always distinguishable from shared ones."""
    import re
    base = re.sub(r"[^a-zA-Z0-9]+", "-", raw.strip().lower()).strip("-")
    if not base:
        base = "data"
    name = f"priv-{base}"
    return name[:63].rstrip("-.")


def add_private(directory: str, collection_label: str | None = None) -> int:
    """Index every supported file under *directory* into the private store.

    Re-running with the same collection upserts (stable chunk IDs), so an
    edited corpus can be refreshed by pointing at the same directory again.
    """
    root = Path(directory).expanduser().resolve()
    if not root.is_dir():
        _err(f"ERROR: not a directory: {root}")
        return 2

    label = collection_label or root.name
    coll_name = _sanitize_collection_name(label)

    code_files = list(_ri.walk_code(root, CODE_EXTENSIONS))
    pdf_files = list(_ri.walk_pdfs(root))
    office_files = list(_walk_office(root))
    if not code_files and not pdf_files and not office_files:
        _err(f"ERROR: no supported files under {root}.")
        _err(f"       Supported: {', '.join(CODE_EXTENSIONS)}, "
             f"{', '.join(OFFICE_EXTENSIONS)}, .pdf")
        return 1

    print(f"Indexing {len(code_files)} text/code + {len(pdf_files)} PDF + "
          f"{len(office_files)} office file(s) from {root}", file=sys.stderr)
    print(f"  -> private collection '{coll_name}' at {PRIVATE_VECTORDB_PATH}",
          file=sys.stderr)

    triples: list[tuple[str, str, dict]] = []
    for p in code_files:
        triples.extend(_ri.process_code_file(p, root, coll_name))
    for p in pdf_files:
        try:
            triples.extend(_ri.process_pdf_file(p, root, coll_name))
        except Exception as e:
            _err(f"  [!] {p}: PDF extraction failed ({e}); skipped")
    for p in office_files:
        triples.extend(_process_office_file(p, root, coll_name))

    if not triples:
        _err("ERROR: nothing extractable (all files empty or unreadable).")
        return 1

    client = _client()
    coll = client.get_or_create_collection(coll_name)
    embedder = _embedder()

    ids = [t[0] for t in triples]
    docs = [t[1] for t in triples]
    metas = [t[2] for t in triples]

    total = len(ids)
    for i in range(0, total, _ri.BATCH_SIZE):
        b_ids = ids[i:i + _ri.BATCH_SIZE]
        b_docs = docs[i:i + _ri.BATCH_SIZE]
        b_metas = metas[i:i + _ri.BATCH_SIZE]
        embs = embedder.encode(b_docs).tolist()
        coll.upsert(ids=b_ids, documents=b_docs, embeddings=embs, metadatas=b_metas)
        print(f"  embedded {min(i + _ri.BATCH_SIZE, total)}/{total} chunks",
              file=sys.stderr)

    sources = _load_sources()
    sources[coll_name] = {
        "label": label,
        "directory": str(root),
        "chunks": total,
        "files": len(code_files) + len(pdf_files) + len(office_files),
    }
    _save_sources(sources)

    print(f"\nDone. Private collection '{coll_name}' now holds {coll.count()} "
          f"chunks and will be retrieved automatically on your next query.",
          file=sys.stderr)
    print("Reminder: retrieved private snippets are sent to whatever client "
          "is connected (including a laptop's VS Code chat history).",
          file=sys.stderr)
    return 0


def list_private() -> int:
    if not os.path.isdir(PRIVATE_VECTORDB_PATH):
        print("No private data indexed yet. Add some with:\n"
              "  ofa --add-private <directory>", file=sys.stderr)
        return 0
    try:
        client = _client()
        collections = client.list_collections()
    except Exception as e:
        _err(f"ERROR: could not open private store: {e}")
        return 1
    if not collections:
        print("No private collections found.", file=sys.stderr)
        return 0
    sources = _load_sources()
    print(f"Private RAG store: {PRIVATE_VECTORDB_PATH}\n", file=sys.stderr)
    for c in collections:
        name = c.name if hasattr(c, "name") else str(c)
        try:
            count = client.get_collection(name).count()
        except Exception:
            count = "?"
        meta = sources.get(name, {})
        src = meta.get("directory", "(unknown source)")
        print(f"  {name}\n    chunks: {count}\n    source: {src}",
              file=sys.stderr)
    return 0


def forget_private(target: str) -> int:
    """Delete one private collection, or everything with ``all``."""
    if not os.path.isdir(PRIVATE_VECTORDB_PATH):
        print("Nothing to forget — no private store exists.", file=sys.stderr)
        return 0
    try:
        client = _client()
    except Exception as e:
        _err(f"ERROR: could not open private store: {e}")
        return 1

    sources = _load_sources()
    existing = [c.name if hasattr(c, "name") else str(c)
                for c in client.list_collections()]

    if target == "all":
        for name in existing:
            try:
                client.delete_collection(name)
            except Exception as e:
                _err(f"  [!] could not delete '{name}': {e}")
        _save_sources({})
        print(f"Deleted all {len(existing)} private collection(s).",
              file=sys.stderr)
        return 0

    # Accept either the raw label the user added or the sanitized name.
    candidate = target if target in existing else _sanitize_collection_name(target)
    if candidate not in existing:
        _err(f"ERROR: no private collection named '{target}'. "
             f"Use 'ofa --list-private' to see what exists.")
        return 1
    try:
        client.delete_collection(candidate)
    except Exception as e:
        _err(f"ERROR: could not delete '{candidate}': {e}")
        return 1
    sources.pop(candidate, None)
    _save_sources(sources)
    print(f"Deleted private collection '{candidate}'.", file=sys.stderr)
    return 0
